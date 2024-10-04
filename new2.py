import json
import pandas as pd
import matplotlib.pyplot as plt
from matplotlib.patches import Patch
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
import os
import subprocess

# Generate pie charts from JSON data
def generate_pie_chart(json_data):
    with open('/home/ec2-user/python/result.json', 'r') as json_file:
        json_content = json.load(json_file)
    data = pd.DataFrame(json_content)

    grouped_data = data.groupby(['Application', 'Services', 'Experiment', 'Results']).size().reset_index(name='count')

    for service in data['Services'].unique():
        service_data = grouped_data[grouped_data['Services'] == service]
        labels = service_data['Experiment'].tolist()
        sizes = service_data['count'].tolist()
        results = service_data['Results'].tolist()
        colors = ['mediumseagreen' if result == 'PASS' else 'tomato' for result in results]

        fig, ax = plt.subplots()
        ax.pie(sizes, labels=labels, colors=colors, startangle=90, wedgeprops={'linewidth': 2, 'edgecolor': 'white'})
        ax.set_title(f'{service}', weight='bold', fontsize=15)

        pass_patch = Patch(color='mediumseagreen', label='Pass')
        fail_patch = Patch(color='tomato', label='Fail')
        plt.legend(handles=[pass_patch, fail_patch], loc='best')
        plt.savefig(f'/home/ec2-user/python/pie_chart_{service}.png')

# Update the Word document with pie charts
def update_doc(json_data):
    with open('/home/ec2-user/python/SRE CHAOS SUMMARY REPORT.json', 'r') as json_file:
        summary_data = json.load(json_file)

    docx_path = '/home/ec2-user/python/SRE-CHAOS-SUMMARY-REPORT-APPNAME.docx'

    if not os.path.exists(docx_path):
        doc = Document()
        doc.add_paragraph('SRE Chaos Summary Report')
        doc.save(docx_path)

    doc = Document(docx_path)

    for i, paragraph in enumerate(doc.paragraphs):
        if 'insert_chart_here' in paragraph.text:
            pie_chart_paragraph = doc.paragraphs[i]
            pie_chart_paragraph.clear()
            for service in summary_data['Services']:
                pie_chart_paragraph.add_run(f'{service}').bold = True
                pie_chart_paragraph.add_run().add_picture(f'/home/ec2-user/python/pie_chart_{service}.png', width=Inches(5))
            doc.paragraphs[i].alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
            break

    p = doc.add_paragraph()
    ref_run = p.add_run('Reference:')
    ref_run.bold = True
    ref_run.font.size = Pt(14)

    p = doc.add_paragraph()
    ref_link = 'https://zap.delta.com/ccoe/docs/operations_process_and_tooling/chaos-engineering/overview/'
    ref_run = p.add_run(ref_link)
    ref_run.italic = True
    ref_run.font.size = Pt(14)
    ref_run.font.color.rgb = RGBColor(0, 0, 255)

    doc.save(docx_path)

# Convert Word to PDF using LibreOffice
def convert_doc_to_pdf(docx_path, pdf_path):
    subprocess.run(['libreoffice', '--headless', '--convert-to', 'pdf', '--outdir', os.path.dirname(pdf_path), docx_path])

# Main script execution
json_detailed_report = "/home/ec2-user/python/result.json"
json_summary_report = "/home/ec2-user/python/SRE CHAOS SUMMARY REPORT.json"

generate_pie_chart(json_detailed_report)
update_doc(json_summary_report)

# Convert Word document to PDF
docx_path = '/home/ec2-user/python/SRE-CHAOS-SUMMARY-REPORT-APPNAME.docx'
pdf_path = '/home/ec2-user/python/SRE-CHAOS-SUMMARY-REPORT-APPNAME.pdf'
convert_doc_to_pdf(docx_path, pdf_path)

