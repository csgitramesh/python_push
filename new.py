
import json
import pandas as pd
import matplotlib.pyplot as plt
from matplotlib.patches import Patch
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
import os

# Generate pie charts from JSON data
def generate_pie_chart(json_data):
    # If json_data is a file path, load the JSON content from the file
    with open('/home/ec2-user/python/result.json', 'r') as json_file:
        json_content = json.load(json_file)
    # Pass the JSON content to read_json
    data = pd.DataFrame(json_content)

    # Group the data by "Services" and "Experiment" columns
    grouped_data = data.groupby(['Application', 'Services', 'Experiment', 'Results']).size().reset_index(name='count')

    # Iterate over each unique service
    for service in data['Services'].unique():
        # Filter data for the current service
        service_data = grouped_data[grouped_data['Services'] == service]

        # Create labels and sizes for the pie chart
        labels = service_data['Experiment'].tolist()
        sizes = service_data['count'].tolist()
        results = service_data['Results'].tolist()

        # Create colors for the pie chart based on "PASS" or "FAIL" status
        colors = ['mediumseagreen' if result == 'PASS' else 'tomato' for result in results]

        # Create a pie chart
        fig, ax = plt.subplots()
        ax.pie(sizes, labels=labels, colors=colors, startangle=90, wedgeprops={'linewidth': 2, 'edgecolor': 'white'})

        # Add title with service name in bold and increased font size
        ax.set_title(f'{service}', weight='bold', fontsize=15)

        # Generate legend
        pass_patch = Patch(color='mediumseagreen', label='Pass')
        fail_patch = Patch(color='tomato', label='Fail')
        plt.legend(handles=[pass_patch, fail_patch], loc='best')

        # Save the pie chart
        plt.savefig(f'/home/ec2-user/python/pie_chart_{service}.png')

# Update the Word document with pie charts
def update_doc(json_data):
    # Load JSON summary report
    with open('/home/ec2-user/python/SRE CHAOS SUMMARY REPORT.json', 'r') as json_file:
        summary_data = json.load(json_file)

    docx_path = '/home/ec2-user/python/SRE-CHAOS-SUMMARY-REPORT-APPNAME.docx'

    # Check if the Word document exists
    if not os.path.exists(docx_path):
        # If it doesn't exist, create a new document
        doc = Document()
        doc.add_paragraph('SRE Chaos Summary Report')
        doc.save(docx_path)

    # Open the Word document
    doc = Document(docx_path)

    # Iterate through paragraphs in the Word document to insert pie charts
    for i, paragraph in enumerate(doc.paragraphs):
        if 'insert_chart_here' in paragraph.text:
            pie_chart_paragraph = doc.paragraphs[i]
            pie_chart_paragraph.clear()

            # Add pie charts for each service
            for service in summary_data['Services']:
                pie_chart_paragraph.add_run(f'{service}').bold = True
                pie_chart_paragraph.add_run().add_picture(f'/home/ec2-user/python/pie_chart_{service}.png', width=Inches(5))
            # Align the paragraph containing the pie charts to the left
            doc.paragraphs[i].alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
            break

    # Add a reference section
    p = doc.add_paragraph()
    ref_run = p.add_run('Reference:')
    ref_run.bold = True
    ref_run.font.size = Pt(14)  # Increase font size for reference section

    # Add a clickable link with blue color
    p = doc.add_paragraph()
    ref_link = 'https://zap.delta.com/ccoe/docs/operations_process_and_tooling/chaos-engineering/overview/'
    ref_run = p.add_run(ref_link)
    ref_run.italic = True
    ref_run.font.size = Pt(14)  # Increase font size for the reference section
    ref_run.font.color.rgb = RGBColor(0, 0, 255)  # Blue color

    # Save the modified document
    doc.save(docx_path)

# Read the JSON files
json_detailed_report = "/home/ec2-user/python/result.json"
json_summary_report = "/home/ec2-user/python/SRE CHAOS SUMMARY REPORT.json"

# Generate pie charts
generate_pie_chart(json_detailed_report)

# Update the Word document
update_doc(json_summary_report)

